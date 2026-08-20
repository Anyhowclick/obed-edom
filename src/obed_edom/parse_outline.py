from __future__ import annotations

import re
from pathlib import Path

from collections import defaultdict

from docx import Document
from docx.oxml.ns import qn

from obed_edom.bible import ABS_REF_RE, REL_VERSE_RE
from obed_edom.models import Cue, OutlineDoc, Paragraph, Run, SlideDraft, TextSpan

SEMANTIC_TAGS = {
    "TITLE",
    "FILLER",
    "FILLER-QR",
    "GIVING-OPTIONS",
    "VERSE",
    "VERSE-CONTINUED",
    "POINT",
    "NUM-POINT",
}
VERSE_TAGS = {"VERSE", "VERSE-CONTINUED"}
GRAPHIC_TAGS = {"TITLE", "FILLER", "FILLER-QR", "GIVING-OPTIONS"}
POINT_TAGS = {"POINT", "NUM-POINT"}
OFFERING_TAGS = {"FILLER-QR", "GIVING-OPTIONS"}

# Longer verse/point tags before their shorter prefixes.
CUE_RE = re.compile(
    r"\[("
    r"FILLER[- ]QR|GIVING[- ]OPTIONS|NUM[- ]POINT|"
    r"VERSE[- ]CONTINUED|VERSE[- ]FROM[- ]PREVIOUS|"
    r"TITLE|FILLER|VERSE|POINT|"
    r"(?:LW|DSK|FW)(?:[-–][^\]]+)?"
    r")\]",
    re.IGNORECASE,
)
POINT_PREFIX_RE = re.compile(r"^(\d+)\.\s+")
STAGE_RE = re.compile(
    r"\[(?:Pray|Instructions|Turn to your neighbours[^\]]*)\]",
    re.IGNORECASE,
)
VERSE_LEAD_RE = re.compile(r"^(\d{1,3})\s+")
GREETING_RE = re.compile(
    r"^(hi church|welcome back|god has great things|won.t you|can you please turn)",
    re.IGNORECASE,
)
DATE_RE = re.compile(
    r"(january|february|march|april|may|june|july|august|september|october|november|december|\d{4})",
    re.IGNORECASE,
)
RHETORIC_RE = re.compile(
    r"^(but,? how |so, |by ourselves|as you prepare|do refer|you may |if you are)",
    re.IGNORECASE,
)

_SEMANTIC_NORMALIZE = {
    "TITLE": "TITLE",
    "FILLER": "FILLER",
    "FILLER-QR": "FILLER-QR",
    "FILLER QR": "FILLER-QR",
    "GIVING-OPTIONS": "GIVING-OPTIONS",
    "GIVING OPTIONS": "GIVING-OPTIONS",
    "VERSE": "VERSE",
    "VERSE-CONTINUED": "VERSE-CONTINUED",
    "VERSE CONTINUED": "VERSE-CONTINUED",
    "VERSE-FROM-PREVIOUS": "VERSE-CONTINUED",
    "VERSE FROM PREVIOUS": "VERSE-CONTINUED",
    "POINT": "POINT",
    "NUM-POINT": "NUM-POINT",
    "NUM POINT": "NUM-POINT",
}


def _is_green(highlight: str | None) -> bool:
    return bool(highlight) and "green" in highlight.lower()


def _highlight_of(run) -> str | None:
    val = run.font.highlight_color
    if val is None:
        return None
    name = getattr(val, "name", None)
    if name and name != "NONE":
        return name.lower()
    xml_val = run._element.find(qn("w:highlight"))
    if xml_val is not None:
        return (xml_val.get(qn("w:val")) or "").lower() or None
    return None


def _color_of(run) -> str | None:
    color = run.font.color
    if color is None or color.rgb is None:
        return None
    return str(color.rgb)


def _is_superscript(run) -> bool:
    return bool(run.font.superscript)


def load_paragraphs(path: Path) -> list[Paragraph]:
    doc = Document(str(path))
    paragraphs: list[Paragraph] = []
    for i, para in enumerate(doc.paragraphs):
        runs: list[Run] = []
        for run in para.runs:
            text = run.text or ""
            if not text:
                continue
            runs.append(
                Run(
                    text=text,
                    bold=bool(run.bold),
                    highlight=_highlight_of(run),
                    superscript=_is_superscript(run),
                    color=_color_of(run),
                )
            )
        paragraphs.append(Paragraph(runs=runs, index=i))
    return paragraphs


def normalize_cue(raw: str) -> Cue:
    inner = raw.strip()[1:-1]
    inner = inner.replace("–", "-").replace("—", "-")
    inner = re.sub(r"\s+", " ", inner).strip().upper()
    semantic_tag = _SEMANTIC_NORMALIZE.get(inner)
    if semantic_tag:
        return Cue(raw=raw, tag=semantic_tag, paragraph=0, offset=0, semantic=True, deck=None)
    if inner.startswith("DSK"):
        deck = "dsk"
    elif inner.startswith("FW"):
        deck = "lw"
    else:
        deck = "lw"
    return Cue(raw=raw, tag=inner, paragraph=0, offset=0, semantic=False, deck=deck)


def _split_paragraph(para: Paragraph) -> list[tuple[str, Cue | None, list[Run]]]:
    """Split a paragraph into (kind, cue_or_none, runs_in_segment) tokens."""
    full = para.text
    tokens: list[tuple[str, Cue | None, list[Run]]] = []
    matches = list(CUE_RE.finditer(full))
    if not matches:
        return [("text", None, para.runs)] if full.strip() else []

    run_map: list[tuple[int, int, Run]] = []
    pos = 0
    for run in para.runs:
        start = pos
        pos += len(run.text)
        run_map.append((start, pos, run))

    def runs_for(a: int, b: int) -> list[Run]:
        out: list[Run] = []
        for start, end, run in run_map:
            if end <= a or start >= b:
                continue
            slice_start = max(0, a - start)
            slice_end = min(len(run.text), b - start)
            chunk = run.text[slice_start:slice_end]
            if chunk:
                out.append(
                    Run(
                        text=chunk,
                        bold=run.bold,
                        highlight=run.highlight,
                        superscript=run.superscript,
                        color=run.color,
                    )
                )
        return out

    cursor = 0
    for match in matches:
        if match.start() > cursor:
            chunk_runs = runs_for(cursor, match.start())
            if "".join(r.text for r in chunk_runs).strip():
                tokens.append(("text", None, chunk_runs))
        cue = normalize_cue(match.group(0))
        cue.paragraph = para.index
        cue.offset = match.start()
        cue.end = match.end()
        tokens.append(("cue", cue, []))
        cursor = match.end()
    if cursor < len(full):
        chunk_runs = runs_for(cursor, len(full))
        if "".join(r.text for r in chunk_runs).strip() or any(r.superscript for r in chunk_runs):
            tokens.append(("text", None, chunk_runs))
    return tokens


def _runs_to_spans(runs: list[Run]) -> list[TextSpan]:
    spans: list[TextSpan] = []
    i = 0
    while i < len(runs):
        run = runs[i]
        if run.superscript and any(ch.isdigit() for ch in run.text):
            number = "".join(ch for ch in run.text if ch.isdigit())
            i += 1
            while i < len(runs) and runs[i].superscript and any(ch.isdigit() for ch in runs[i].text):
                number += "".join(ch for ch in runs[i].text if ch.isdigit())
                i += 1
            if not number:
                continue
            rest_parts: list[Run] = []
            while i < len(runs) and not (
                runs[i].superscript and any(ch.isdigit() for ch in runs[i].text)
            ):
                rest_parts.append(runs[i])
                i += 1
            if spans and spans[-1].text and not spans[-1].text[-1].isspace():
                spans.append(TextSpan(text=" "))
            spans.append(TextSpan(text=number, verse_number=number))
            if rest_parts:
                if rest_parts[0].text and not rest_parts[0].text[0].isspace():
                    spans.append(TextSpan(text=" "))
                for r in rest_parts:
                    t = r.text
                    if STAGE_RE.search(t):
                        t = STAGE_RE.sub("", t)
                    if t:
                        spans.append(
                            TextSpan(
                                text=t,
                                bold=bool(r.bold),
                                highlight=r.highlight,
                            )
                        )
            continue
        text = run.text
        if STAGE_RE.search(text):
            text = STAGE_RE.sub("", text)
        if text:
            spans.append(
                TextSpan(
                    text=text,
                    bold=run.bold,
                    highlight=run.highlight,
                )
            )
        i += 1
    return spans


def _green_title(runs: list[Run]) -> str:
    bits = [r.text for r in runs if _is_green(r.highlight)]
    return "".join(bits).strip()


def _is_verse_runs(runs: list[Run]) -> bool:
    if any(r.superscript and r.text.strip().isdigit() for r in runs):
        return True
    text = "".join(r.text for r in runs).strip()
    return bool(VERSE_LEAD_RE.match(text))


def _is_bare_reference(text: str) -> bool:
    stripped = STAGE_RE.sub("", text).strip()
    if len(stripped) > 48:
        return False
    match = ABS_REF_RE.search(stripped)
    return bool(match) and match.group(0).strip() in stripped and len(stripped) - len(match.group(0)) < 12


def _is_speaker_commentary(text: str, runs: list[Run]) -> bool:
    stripped = STAGE_RE.sub("", text).strip()
    if not stripped:
        return True
    if _is_verse_runs(runs):
        return False
    if any(_is_green(r.highlight) for r in runs):
        return False
    if stripped.lower() in {"conclusion", "introduction", "closing", "altar call"}:
        return True
    if GREETING_RE.match(stripped):
        return True
    if RHETORIC_RE.match(stripped):
        return True
    if stripped.endswith("?") and len(stripped) > 40:
        return True
    if len(stripped) > 180 and not _is_verse_runs(runs):
        return True
    return False


def _detect_header(paragraphs: list[Paragraph]) -> tuple[str, str, str]:
    date_line = ""
    titles: list[str] = []
    for para in paragraphs:
        text = STAGE_RE.sub("", para.text).strip()
        if not text:
            if titles:
                break
            continue
        if CUE_RE.search(text):
            break
        if GREETING_RE.match(text):
            break
        if DATE_RE.search(text) and ("church" in text.lower() or re.search(r"\d{4}", text)):
            date_line = text
            continue
        bold = any(r.bold for r in para.runs) or len(text) < 80
        if bold:
            titles.append(text.rstrip(":").strip())
        else:
            break
    series_title = titles[0] if titles else ""
    series_subtitle = " ".join(titles[1:]) if len(titles) > 1 else ""
    return series_title, series_subtitle, date_line


def _prev_relative_verse(paragraphs: list[Paragraph], index: int) -> bool:
    for j in range(index - 1, -1, -1):
        text = STAGE_RE.sub("", paragraphs[j].text).strip()
        if not text:
            continue
        return bool(REL_VERSE_RE.search(text))
    return False


def _new_draft(
    cue: Cue,
    paragraphs: list[Paragraph],
    point_number: int | None = None,
) -> SlideDraft:
    return SlideDraft(
        cue_tag=cue.tag,
        cue_raw=cue.raw,
        cue_paragraph=cue.paragraph,
        cue_offset=cue.offset,
        cue_end=cue.end,
        source_paragraphs=[cue.paragraph],
        force_verse=_prev_relative_verse(paragraphs, cue.paragraph) or cue.tag in VERSE_TAGS,
        point_number=point_number,
    )


class ListNumberResolver:
    """Resolve Word auto-numbering values (e.g. outline point 9.)."""

    def __init__(self, doc: Document):
        self.num_to_abstract: dict[str, str] = {}
        self.abstract_levels: dict[str, dict[str, int]] = {}
        self.num_level_overrides: dict[tuple[str, str], int] = {}
        self.counters: dict[tuple[str, str], int] = defaultdict(int)
        try:
            numbering = doc.part.numbering_part.element
        except Exception:  # noqa: BLE001
            return

        for num in numbering.findall(qn("w:num")):
            num_id = num.get(qn("w:numId"))
            if not num_id:
                continue
            abstract = num.find(qn("w:abstractNumId"))
            if abstract is not None:
                self.num_to_abstract[num_id] = abstract.get(qn("w:val")) or ""
            for override in num.findall(qn("w:lvlOverride")):
                ilvl = override.get(qn("w:ilvl")) or "0"
                start = override.find(qn("w:startOverride"))
                if start is not None:
                    self.num_level_overrides[(num_id, ilvl)] = int(start.get(qn("w:val")) or 1)

        for abstract in numbering.findall(qn("w:abstractNum")):
            abs_id = abstract.get(qn("w:abstractNumId"))
            if not abs_id:
                continue
            levels: dict[str, int] = {}
            for lvl in abstract.findall(qn("w:lvl")):
                ilvl = lvl.get(qn("w:ilvl")) or "0"
                start = lvl.find(qn("w:start"))
                levels[ilvl] = int(start.get(qn("w:val")) if start is not None else 1)
            self.abstract_levels[abs_id] = levels

    def _start_for(self, num_id: str, ilvl: str) -> int:
        override = self.num_level_overrides.get((num_id, ilvl))
        if override is not None:
            return override
        abs_id = self.num_to_abstract.get(num_id, "")
        return self.abstract_levels.get(abs_id, {}).get(ilvl, 1)

    def number_for_paragraph(self, para) -> int | None:
        p_pr = para._element.find(qn("w:pPr"))
        if p_pr is None:
            return None
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            return None
        num_id_el = num_pr.find(qn("w:numId"))
        if num_id_el is None:
            return None
        num_id = num_id_el.get(qn("w:val"))
        if not num_id:
            return None
        ilvl_el = num_pr.find(qn("w:ilvl"))
        ilvl = ilvl_el.get(qn("w:val")) if ilvl_el is not None else "0"
        key = (num_id, ilvl)
        self.counters[key] += 1
        return self._start_for(num_id, ilvl) + self.counters[key] - 1


def _literal_point_number(text: str) -> int | None:
    stripped = STAGE_RE.sub("", text).strip()
    match = POINT_PREFIX_RE.match(stripped)
    if match:
        return int(match.group(1))
    return None


def _mark_verse_follows(blocks: list[SlideDraft]) -> None:
    for i, draft in enumerate(blocks):
        if draft.cue_tag not in POINT_TAGS:
            continue
        if i + 1 < len(blocks) and blocks[i + 1].cue_tag in VERSE_TAGS:
            draft.verse_follows = True
            draft.following_verse_index = i + 1


def parse_outline(path: Path | str) -> OutlineDoc:
    path = Path(path)
    doc = Document(str(path))
    paragraphs = load_paragraphs(path)
    list_numbers = ListNumberResolver(doc)
    series_title, series_subtitle, date_line = _detect_header(paragraphs)
    full_text = "\n".join(p.text for p in paragraphs)

    para_numbers: dict[int, int] = {}
    for i, para in enumerate(doc.paragraphs):
        number = list_numbers.number_for_paragraph(para)
        if number is None:
            number = _literal_point_number(para.text)
        if number is not None:
            para_numbers[i] = number

    current: SlideDraft | None = None
    blocks: list[SlideDraft] = []
    saw_offering = False

    def close_current() -> None:
        nonlocal current
        if current is not None:
            blocks.append(current)
            current = None

    def append_body(draft: SlideDraft | None, runs: list[Run], para_index: int) -> None:
        if draft is None:
            return
        spans = _runs_to_spans(runs)
        green = _green_title(runs)
        if green:
            draft.green_title = green
        if spans:
            if (
                draft.body_spans
                and draft.body_spans[-1].text
                and not draft.body_spans[-1].text[-1].isspace()
                and spans[0].text
                and not spans[0].text[0].isspace()
            ):
                draft.body_spans.append(TextSpan(text=" "))
            draft.body_spans.extend(spans)
            if para_index not in draft.source_paragraphs:
                draft.source_paragraphs.append(para_index)

    def treat_as_notes(draft: SlideDraft, runs: list[Run]) -> bool:
        if draft.cue_tag in GRAPHIC_TAGS:
            return True
        if draft.cue_tag in POINT_TAGS and draft.body and not _is_verse_runs(runs):
            return True
        if draft.has_verse_numbers and not _is_verse_runs(runs):
            return True
        return False

    for para in paragraphs:
        tokens = _split_paragraph(para)
        if not tokens:
            continue

        has_cue = any(kind == "cue" for kind, _, _ in tokens)
        if has_cue:
            for kind, cue, runs in tokens:
                if kind == "cue" and cue is not None:
                    if cue.tag in OFFERING_TAGS or "OFFERING" in cue.tag or "GIVING" in cue.tag:
                        saw_offering = True
                    if cue.semantic or cue.tag in SEMANTIC_TAGS:
                        close_current()
                        point_number = para_numbers.get(para.index) if cue.tag == "NUM-POINT" else None
                        current = _new_draft(cue, paragraphs, point_number=point_number)
                    elif cue.deck == "lw":
                        close_current()
                        current = _new_draft(cue, paragraphs, point_number=para_numbers.get(para.index))
                    else:
                        # Legacy DSK-only cue: still a content block if no semantic outline.
                        close_current()
                        current = _new_draft(cue, paragraphs, point_number=para_numbers.get(para.index))
                elif kind == "text":
                    append_body(current, runs, para.index)
            continue

        text = STAGE_RE.sub("", para.text).strip()
        if not text:
            continue
        if current is None:
            continue
        if _is_bare_reference(para.text) and current.cue_tag not in VERSE_TAGS:
            current.speaker_notes.append(text)
            continue
        if _is_speaker_commentary(para.text, para.runs) and not _is_verse_runs(para.runs):
            current.speaker_notes.append(text)
            continue
        if treat_as_notes(current, para.runs):
            current.speaker_notes.append(text)
            continue
        append_body(current, para.runs, para.index)

    close_current()
    _mark_verse_follows(blocks)

    context = "offering" if saw_offering else "sermon"
    return OutlineDoc(
        path=path,
        paragraphs=paragraphs,
        series_title=series_title,
        series_subtitle=series_subtitle,
        date_line=date_line,
        context=context,
        blocks=blocks,
        lw_slides=list(blocks),
        dsk_slides=list(blocks),
        full_text=full_text,
    )
