from __future__ import annotations

import re
import shutil
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from sermon_slides.models import OutlineDoc, SlideSpec
from sermon_slides.parse_outline import CUE_RE, SEMANTIC_TAGS, normalize_cue

# Word ST_Highlight has no "turquoise"; the UI colour Turquoise is stored as cyan.
LW_HIGHLIGHT = "cyan"
DSK_HIGHLIGHT = "yellow"


@dataclass
class _Op:
    start: int
    end: int
    tags: list[tuple[str, str]]  # (text, highlight)
    order: int = 0


def _make_run(text: str, highlight: str, bold: bool = True):
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if bold:
        rpr.append(OxmlElement("w:b"))
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), highlight)
    rpr.append(hl)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _run_map(paragraph) -> list[tuple[int, int, object]]:
    mapping: list[tuple[int, int, object]] = []
    pos = 0
    for run in paragraph.runs:
        text = run.text or ""
        start = pos
        pos += len(text)
        mapping.append((start, pos, run))
    return mapping


def _split_run_clean(run, local_index: int) -> None:
    text = run.text or ""
    if local_index <= 0 or local_index >= len(text):
        return
    tail = text[local_index:]
    run.text = text[:local_index]
    new_r = OxmlElement("w:r")
    rpr = run._element.find(qn("w:rPr"))
    if rpr is not None:
        new_r.append(deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = tail
    new_r.append(t)
    run._element.addnext(new_r)


def _apply_ops(paragraph, ops: list[_Op]) -> None:
    """Apply delete/insert ops using original coordinates, right to left."""
    if not ops:
        return
    ordered = sorted(ops, key=lambda o: (o.start, o.order), reverse=True)
    for op in ordered:
        # Refresh run map each time because previous splits changed the tree.
        mapping = _run_map(paragraph)
        if not mapping and op.start == 0 and op.end == 0:
            for text, hl in reversed(op.tags):
                paragraph._p.append(_make_run(text, hl))
            continue

        def run_at(index: int, at_end: bool = False):
            if not mapping:
                return None, 0
            if at_end:
                for start, end, run in mapping:
                    if start <= index <= end:
                        return run, index - start
            for start, end, run in mapping:
                if start <= index < end:
                    return run, index - start
            start, end, run = mapping[-1]
            return run, len(run.text or "")

        if op.end > op.start:
            start_run, start_local = run_at(op.start)
            end_run, end_local = run_at(op.end, at_end=True)
            if start_run is None:
                continue
            _split_run_clean(start_run, start_local)
            mapping = _run_map(paragraph)
            end_run, end_local = run_at(op.end, at_end=True)
            if end_run is not None:
                _split_run_clean(end_run, end_local)
            mapping = _run_map(paragraph)
            to_delete = []
            for start, end, run in mapping:
                if start >= op.start and end <= op.end and start < op.end:
                    to_delete.append(run)
            anchor = None
            for start, end, run in mapping:
                if end <= op.start:
                    anchor = run
            for run in to_delete:
                parent = run._element.getparent()
                if parent is not None:
                    parent.remove(run._element)
            insert_after = anchor._element if anchor is not None else None
            _insert_runs(paragraph, insert_after, op.tags)
        else:
            # Pure insert at op.start. Leave a space before the original text when needed.
            tags = list(op.tags)
            if tags:
                last_text, last_hl = tags[-1]
                if not last_text.endswith(" "):
                    tags[-1] = (last_text + " ", last_hl)
            start_run, start_local = run_at(op.start)
            if start_run is None:
                _insert_runs(paragraph, None, tags)
                continue
            if start_local > 0:
                _split_run_clean(start_run, start_local)
                mapping = _run_map(paragraph)
                insert_after = start_run._element
            elif op.start == 0:
                insert_after = None
            else:
                prev = None
                for start, end, run in mapping:
                    if start < op.start:
                        prev = run
                insert_after = prev._element if prev is not None else None
            _insert_runs(paragraph, insert_after, tags)


def _insert_runs(paragraph, insert_after, tags: list[tuple[str, str]]):
    parent = paragraph._p
    for text, hl in tags:
        el = _make_run(text, hl)
        if insert_after is not None:
            insert_after.addnext(el)
        else:
            ppr = parent.find(qn("w:pPr"))
            if ppr is not None:
                ppr.addnext(el)
            else:
                parent.insert(0, el)
        insert_after = el
    return insert_after


def _operator_label(tag: str) -> str:
    return f"[{tag}]"


def _highlight_for(tag: str) -> str:
    if tag.upper().startswith("DSK"):
        return DSK_HIGHLIGHT
    return LW_HIGHLIGHT


def _cue_ops_for_block(draft, specs: list[SlideSpec]) -> _Op | None:
    cue_specs = [s for s in specs if s.bind == "cue"]
    if not cue_specs:
        # Still delete the semantic cue (e.g. [VERSE] whose slides bind to the body).
        if draft.cue_tag in SEMANTIC_TAGS:
            return _Op(start=draft.cue_offset, end=draft.cue_end, tags=[], order=0)
        return None
    tags: list[tuple[str, str]] = []
    seen: list[str] = []
    # Keep deck order: LW then DSK, PRE before anything else at the cue.
    ordered = sorted(
        cue_specs,
        key=lambda s: (0 if s.deck == "lw" else 1, 0 if s.role == "pre" else 1, s.chunk_index),
    )
    for spec in ordered:
        label = _operator_label(spec.operator_tag or spec.cue_tag)
        if label in seen:
            continue
        seen.append(label)
        tags.append((label, _highlight_for(spec.operator_tag or spec.cue_tag)))
    return _Op(start=draft.cue_offset, end=draft.cue_end, tags=tags, order=0)


def _chunk_search_needle(spec: SlideSpec) -> str:
    """Distinctive leading text of a verse chunk, without a leading verse number."""
    body = re.sub(r"\s+", " ", (spec.body or "").replace("\xa0", " ")).strip()
    body = re.sub(r"^\d+\s*", "", body)
    body = re.sub(r"^[.…\s]+", "", body).strip()
    if len(body) < 8:
        return ""
    return body[:28]


def _find_text_offset(
    outline: OutlineDoc,
    para_indices: list[int],
    needle: str,
    cue_spans: dict[int, tuple[int, int]],
) -> tuple[int, int] | None:
    if not needle:
        return None
    for pidx in para_indices:
        if pidx < 0 or pidx >= len(outline.paragraphs):
            continue
        text = outline.paragraphs[pidx].text.replace("\xa0", " ")
        start = cue_spans[pidx][1] if pidx in cue_spans else 0
        idx = text.find(needle, start)
        if idx < 0:
            idx = text.find(needle)
        if idx >= 0:
            return pidx, idx
    return None


def _locate_verse_chunk(
    outline: OutlineDoc,
    spec: SlideSpec,
    spots: list[tuple[int, int, str]],
    deck_used: set[tuple[int, int]],
    cue_spans: dict[int, tuple[int, int]],
) -> tuple[int, int] | None:
    """Offset where this verse *slide* starts — not the first verse of the whole passage."""
    if spec.anchor_verse:
        for pidx, off, number in spots:
            if number == spec.anchor_verse and (pidx, off) not in deck_used:
                return pidx, off
    loc = _find_text_offset(
        outline, spec.source_paragraphs or [], _chunk_search_needle(spec), cue_spans
    )
    if loc is not None:
        return loc
    for pidx, off, _number in spots:
        if (pidx, off) not in deck_used:
            return pidx, off
    if spots:
        return spots[0][0], spots[0][1]
    paras = spec.source_paragraphs or [0]
    for pidx in paras:
        if pidx in cue_spans:
            return pidx, cue_spans[pidx][1]
        text = outline.paragraphs[pidx].text if pidx < len(outline.paragraphs) else ""
        if any(ch.isdigit() for ch in text) and "VERSE" not in text.upper():
            return pidx, 0
    return paras[-1], 0


def _verse_positions(outline: OutlineDoc, para_indices: list[int]) -> list[tuple[int, int, str]]:
    """Return (para_index, offset, verse_number) for superscript verse numbers."""
    wanted = set(para_indices)
    found: list[tuple[int, int, str]] = []
    for para in outline.paragraphs:
        if para.index not in wanted:
            continue
        offset = 0
        i = 0
        runs = para.runs
        while i < len(runs):
            run = runs[i]
            text = run.text or ""
            if run.superscript and any(ch.isdigit() for ch in text):
                digit_at = next(j for j, ch in enumerate(text) if ch.isdigit())
                start = offset + digit_at
                number = "".join(ch for ch in text if ch.isdigit())
                offset += len(text)
                i += 1
                while i < len(runs) and runs[i].superscript and any(
                    ch.isdigit() for ch in (runs[i].text or "")
                ):
                    number += "".join(ch for ch in runs[i].text if ch.isdigit())
                    offset += len(runs[i].text or "")
                    i += 1
                found.append((para.index, start, number))
                continue
            offset += len(text)
            i += 1
    return found


def _body_insert_ops(
    outline: OutlineDoc,
    specs: list[SlideSpec],
    cue_spans: dict[int, tuple[int, int]] | None = None,
) -> dict[int, list[_Op]]:
    """Insert operator tags at each chunk start in verse body paragraphs."""
    by_para: dict[int, list[_Op]] = {}
    # POST uses the same [LW]/[DSK-PP] labels as the following verse slides.
    # Putting those on the verse line doubles cues; PRE already marks the point.
    body_specs = [s for s in specs if s.bind == "verse_body" and s.role == "verse"]
    if not body_specs:
        return by_para
    cue_spans = cue_spans or {}

    positions_cache: dict[tuple[int, ...], list[tuple[int, int, str]]] = {}

    def positions_for(source: list[int]) -> list[tuple[int, int, str]]:
        key = tuple(source)
        if key not in positions_cache:
            positions_cache[key] = _verse_positions(outline, source)
        return positions_cache[key]

    grouped: dict[tuple[int, int], list[SlideSpec]] = {}
    used: dict[str, set[tuple[int, int]]] = {"lw": set(), "dsk": set()}
    ordered = sorted(
        body_specs,
        key=lambda s: (s.block_index, s.chunk_index, 0 if s.deck == "lw" else 1),
    )
    for spec in ordered:
        spots = positions_for(spec.source_paragraphs or [0])
        deck_used = used.setdefault(spec.deck, set())
        loc = _locate_verse_chunk(outline, spec, spots, deck_used, cue_spans)
        if loc is None:
            continue
        para_index, offset = loc
        if para_index in cue_spans:
            _c0, cue_end = cue_spans[para_index]
            if offset < cue_end:
                offset = cue_end
        deck_used.add((para_index, offset))
        grouped.setdefault((para_index, offset), []).append(spec)

    for (para_index, offset), group in grouped.items():
        group.sort(key=lambda s: (s.chunk_index, 0 if s.deck == "lw" else 1))
        tags: list[tuple[str, str]] = []
        seen: list[str] = []
        for spec in group:
            label = _operator_label(spec.operator_tag or spec.cue_tag)
            if label in seen:
                continue
            seen.append(label)
            tags.append((label, _highlight_for(spec.operator_tag or spec.cue_tag)))
        by_para.setdefault(para_index, []).append(_Op(start=offset, end=offset, tags=tags, order=1))
    return by_para


def annotate_outline(
    outline: OutlineDoc,
    lw: list[SlideSpec],
    dsk: list[SlideSpec],
    dest: Path,
) -> Path:
    """Copy the source outline and replace semantic cues with operator [LW]/[DSK] tags."""
    dest = Path(dest)
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(outline.path, dest)
    doc = Document(str(dest))
    all_specs = list(lw) + list(dsk)
    by_block: dict[int, list[SlideSpec]] = {}
    for spec in all_specs:
        by_block.setdefault(spec.block_index, []).append(spec)

    para_ops: dict[int, list[_Op]] = {}
    for i, draft in enumerate(outline.blocks):
        specs = by_block.get(i, [])
        op = _cue_ops_for_block(draft, specs)
        if op is not None:
            para_ops.setdefault(draft.cue_paragraph, []).append(op)

    cue_spans = {draft.cue_paragraph: (draft.cue_offset, draft.cue_end) for draft in outline.blocks}
    body_ops = _body_insert_ops(outline, all_specs, cue_spans)
    for para_index, ops in body_ops.items():
        para_ops.setdefault(para_index, []).extend(ops)

    for para_index, ops in para_ops.items():
        if para_index < 0 or para_index >= len(doc.paragraphs):
            continue
        merged: list[_Op] = []
        inserts: dict[int, _Op] = {}
        for op in ops:
            if op.end > op.start:
                merged.append(op)
                continue
            existing = inserts.get(op.start)
            if existing is None:
                inserts[op.start] = op
                merged.append(op)
            else:
                have = {text for text, _hl in existing.tags}
                for text, hl in op.tags:
                    if text not in have:
                        existing.tags.append((text, hl))
                        have.add(text)
        _apply_ops(doc.paragraphs[para_index], merged)

    doc.save(str(dest))
    return dest


def extract_operator_cues(path: Path | str) -> list[str]:
    """Flatten operator cue tags in document order (for tests / golden checks)."""
    doc = Document(str(path))
    tags: list[str] = []
    for para in doc.paragraphs:
        for match in CUE_RE.finditer(para.text):
            cue = normalize_cue(match.group(0))
            if cue.semantic:
                continue
            tags.append(cue.tag)
    return tags
