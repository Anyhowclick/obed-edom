"""Outline fixtures shared by the parse and outline-check tests.

The real operator outlines live in a gitignored folder, so anything that has to
run everywhere builds its own `.docx` here. `verse_after_point_variant` exists
because `Sermon BC.docx` predates `[VERSE-AFTER-POINT]`: it still writes a plain
`[VERSE]` after a point, which under the current grammar is a verse-only slide.
Retagging those cues is what asks for the point-plus-verse masters again.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

from obed_edom.parse_outline import CUE_RE, POINT_TAGS, parse_outline


def _replace_span(paragraph, start: int, end: int, text: str) -> None:
    """Swap the [start, end) slice of a paragraph's text, keeping other runs."""
    pos = 0
    written = False
    for run in paragraph.runs:
        run_start, run_end = pos, pos + len(run.text)
        pos = run_end
        if run_end <= start or run_start >= end:
            continue
        head = run.text[: max(0, start - run_start)]
        tail = run.text[max(0, min(len(run.text), end - run_start)) :]
        run.text = head + ("" if written else text) + tail
        written = True


def verse_after_point_variant(src: Path | str, dest: Path | str) -> Path:
    """Copy an outline, retagging every `[VERSE]` that directly follows a point."""
    src, dest = Path(src), Path(dest)
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dest)

    outline = parse_outline(src)
    targets = [
        outline.blocks[i + 1]
        for i, block in enumerate(outline.blocks[:-1])
        if block.cue_tag in POINT_TAGS and outline.blocks[i + 1].cue_tag == "VERSE"
    ]
    if not targets:
        return dest

    doc = Document(str(dest))
    # Late paragraphs first so earlier offsets stay valid as spans change length.
    for block in sorted(targets, key=lambda b: (b.cue_paragraph, b.cue_offset), reverse=True):
        if block.cue_paragraph >= len(doc.paragraphs):
            continue
        _replace_span(
            doc.paragraphs[block.cue_paragraph],
            block.cue_offset,
            block.cue_end,
            "[VERSE-AFTER-POINT]",
        )
    doc.save(str(dest))
    return dest


def duplicate_as_plain_verse(path: Path | str, needle: str) -> Path:
    """Re-cue a copy of the paragraph holding `needle` as a plain `[VERSE]`.

    Gives the passage a point-plus-verse slide and a verse-only slide, which is
    how the same verse box lands on two slides now that nothing is implicit.
    """
    import copy

    path = Path(path)
    doc = Document(str(path))
    for para in doc.paragraphs:
        if needle not in para.text:
            continue
        match = CUE_RE.search(para.text)
        if not match:
            continue
        clone = copy.deepcopy(para._p)
        para._p.addnext(clone)
        from docx.text.paragraph import Paragraph as _Paragraph

        _replace_span(_Paragraph(clone, para._parent), match.start(), match.end(), "[VERSE]")
        break
    doc.save(str(path))
    return path


def build_outline(dest: Path | str, lines, *, title: str = "Passion for God (Part 4): ") -> Path:
    """Build a minimal outline. A line is a string, or (text, {run options})."""
    dest = Path(dest)
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph(title)
    doc.add_paragraph("Truth and Contentment")
    doc.add_paragraph("")
    for line in lines:
        para = doc.add_paragraph()
        parts = [(line, {})] if isinstance(line, str) else line
        for text, opts in parts:
            run = para.add_run(text)
            if opts.get("superscript"):
                run.font.superscript = True
            if opts.get("bold"):
                run.bold = True
    doc.save(str(dest))
    return dest


def operator_cue_counts(path: Path | str) -> tuple[int, int]:
    """(LW, DSK) operator cue counts in document order."""
    lw = dsk = 0
    for para in Document(str(path)).paragraphs:
        for match in CUE_RE.finditer(para.text):
            inner = match.group(1).upper()
            if inner.startswith(("LW", "FW")):
                lw += 1
            elif inner.startswith("DSK"):
                dsk += 1
    return lw, dsk
