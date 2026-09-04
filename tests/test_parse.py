"""Parse and generate-pass tests.

Scripting cannot write superscript: pass 2 Copy Style / Paste Style from the seed verse number.
Pass 1 leaves the deck open; pass 2 styles then exports. Do not export in pass 1.
"""

from dataclasses import replace
from pathlib import Path
from unittest.mock import patch

import pytest
from docx.oxml.ns import qn
from obed_edom import keynote_app
from obed_edom.annotate import annotate_outline, extract_operator_cues
from obed_edom.bible import _parse_gateway_html, check_bible, fetch_passage
from obed_edom.models import StyledRun
from obed_edom.parse_outline import parse_outline
from obed_edom.slide_map import _split_styled_runs, map_slides
from outline_fixtures import (
    build_outline,
    duplicate_as_plain_verse,
    operator_cue_counts,
    verse_after_point_variant,
)

ROOT = Path(__file__).resolve().parents[1]
OUTLINES = ROOT / "Sermon Outlines"
_NEED = ["Sermon BC.docx", "Offering JX.docx", "Offering JX_EXPECTED_RESULT.docx"]
pytestmark = pytest.mark.skipif(
    not all((OUTLINES / name).is_file() for name in _NEED),
    reason="Sermon Outlines/ fixtures are local operator files (gitignored)",
)


def test_sermon_cues():
    outline = parse_outline(OUTLINES / "Sermon BC.docx")
    assert outline.context == "sermon"
    assert "Passion" in outline.series_title
    tags = [s.cue_tag for s in outline.blocks]
    assert tags[0] == "TITLE"
    assert "VERSE" in tags
    assert "NUM-POINT" in tags
    assert "POINT" in tags
    truthful = next(s for s in outline.blocks if "Truthful" in (s.green_title or s.body))
    assert truthful.cue_tag == "NUM-POINT"
    assert truthful.point_number == 1
    examine = next(s for s in outline.blocks if "Examine" in s.body)
    assert examine.cue_tag == "POINT"
    communion = next(s for s in outline.blocks if "Communion" in s.body)
    assert communion.cue_tag == "POINT"
    v27 = next(s for s in outline.blocks if s.cue_tag == "VERSE" and "27" in s.body)
    assert "Spirit" in v27.body
    # A plain [VERSE] after a point is a verse-only slide, so the point stays a
    # static PRE. Only [VERSE-AFTER-POINT] Magic Moves out of it.
    assert not any(s.verse_follows for s in outline.blocks)


def test_verse_after_point_links_to_its_point(tmp_path):
    variant = verse_after_point_variant(OUTLINES / "Sermon BC.docx", tmp_path / "BC_VAP.docx")
    outline = parse_outline(variant)
    truthful = next(s for s in outline.blocks if "Truthful" in (s.green_title or s.body))
    assert truthful.verse_follows
    assert outline.blocks[truthful.following_verse_index].cue_tag == "VERSE-AFTER-POINT"
    # Examine is followed by another point, so nothing moves out of it.
    examine = next(s for s in outline.blocks if "Examine" in s.body)
    assert not examine.verse_follows


@pytest.mark.parametrize("cue", ["[VERSE-AFTER-POINT]", "[VERSE AFTER POINT]"])
def test_verse_after_point_spelling(tmp_path, cue):
    path = build_outline(
        tmp_path / f"{cue.strip('[]').replace(' ', '_')}.docx",
        ["[NUM-POINT] Be Truthful in Love", f"{cue} Ezekiel 36:27 And I will put my Spirit in you."],
    )
    assert [b.cue_tag for b in parse_outline(path).blocks] == ["NUM-POINT", "VERSE-AFTER-POINT"]


def test_offering_cues():
    outline = parse_outline(OUTLINES / "Offering JX.docx")
    assert outline.context == "offering"
    tags = [s.cue_tag for s in outline.blocks]
    assert tags[0] == "FILLER"
    assert "FILLER-QR" in tags
    assert "GIVING-OPTIONS" in tags
    assert tags.count("VERSE") >= 1
    assert "VERSE" in tags or "VERSE-CONTINUED" in tags
    verse_like = [s for s in outline.blocks if s.cue_tag in {"VERSE", "VERSE-CONTINUED"}]
    assert len(verse_like) >= 2
    verse = next(s for s in outline.blocks if s.has_verse_numbers)
    assert "31" in verse.body and "33" in verse.body


def test_mapping_masters():
    sermon = parse_outline(OUTLINES / "Sermon BC.docx")
    lw, dsk, _ = map_slides(sermon)
    assert lw[0].master == "TITLE"
    assert lw[0].cue_tag == "LW-TITLE"
    assert all(s.cue_tag != "LW-TITLE" for s in dsk)

    truthful_pre = next(
        s
        for s in lw
        if s.master == "NUMBERED POINT PRE" and "Truthful" in s.text_items.get(1, "")
    )
    assert truthful_pre.text_items.get(4) == "1"
    title_runs = truthful_pre.styled_items.get(1) or []
    styles = {r.style: r.text for r in title_runs}
    assert any(r.style == "normal" and "Be" in r.text for r in title_runs)
    assert any(r.style == "highlight" and "Truthful" in r.text for r in title_runs)
    assert any(r.style == "normal" and "Love" in r.text for r in title_runs)
    assert truthful_pre.item_palettes.get(1) == "lw_point"
    pre_size = truthful_pre.text_item_font_sizes.get(1, 180)
    assert pre_size >= 170

    # This outline writes a plain [VERSE] after each point, so there is nothing
    # to Magic Move into and no point-plus-verse slide.
    assert truthful_pre.transition is None
    assert not any(s.role == "post" for s in lw)
    assert not any(s.role == "post" for s in dsk)

    communion_pre = next(
        s for s in lw if s.master == "NON-NUMBERED POINT PRE" and "Communion" in (s.body or "")
    )
    pre_text = communion_pre.text_items.get(1, "")
    assert "Powerful Presence" in pre_text
    assert pre_text.count("\n") == 1
    assert pre_text.split("\n")[0].endswith("Communion")
    assert not (communion_pre.text_items.get(2) or "").strip()
    assert not (communion_pre.text_items.get(3) or "").strip()

    assert any(s.master == "VERSES" for s in lw)
    v26 = next(s for s in lw if s.is_verse and s.role == "verse" and "new heart" in s.body)
    styles = [r.style for runs in v26.styled_items.values() for r in runs]
    texts = [r.text for runs in v26.styled_items.values() for r in runs]
    assert "verse_number" in styles
    assert any(t.strip() == "26" for t in texts)
    assert not v26.body.startswith("Ezekiel")
    assert any(r.style == "highlight" and "new heart" in r.text for runs in v26.styled_items.values() for r in runs)

    dsk_pre = next(s for s in dsk if s.master == "Num Point with Verse-Pre")
    assert dsk_pre.transition is None
    v27 = next(s for s in dsk if s.is_verse and s.role == "verse" and "keep my laws" in s.body)
    assert v27.master == "Verse Standard (Variation 2)"
    assert "NIV" not in v27.header
    assert "Ezekiel 36" in v27.header

    offering = parse_outline(OUTLINES / "Offering JX.docx")
    lw, dsk, flags = map_slides(offering)
    assert any(s.master == "BLANK" and s.cue_tag == "LW-OFFERING FILLER" for s in lw)
    assert sum(1 for s in lw if s.cue_tag == "LW-OFFERING FILLER") == 2
    assert any(s.master == "BLANK" and s.cue_tag == "LW-GIVING OPTIONS" for s in lw)
    assert any("QR" in s.master for s in dsk)
    assert any(s.master.startswith("Verse (Offering") for s in dsk)
    assert any("giving" in f.message.lower() for f in flags)
    lw_verses = [s for s in lw if s.is_verse and s.role == "verse"]
    assert len(lw_verses) >= 3
    assert all(s.text_item_heights.get(2, 0) >= 550 for s in lw_verses)
    assert all(len(s.body) <= 320 for s in lw_verses)
    combined = next((s for s in lw_verses if "31" in s.body and "32" in s.body), None)
    assert combined is not None
    solo_33 = next((s for s in lw_verses if s.body.strip().startswith("33")), None)
    assert solo_33 is not None
    dsk_verses = [s for s in dsk if s.is_verse and s.role == "verse"]
    assert dsk_verses
    assert all("(MSG)" in s.header for s in dsk_verses)
    assert all(s.text_item_widths.get(1) == 1540 for s in dsk_verses)
    assert all(
        len(s.body) <= 190 for s in dsk_verses if s.semantic_tag != "VERSE-CONTINUED"
    )
    assert min(len(s.body) for s in dsk_verses) >= 40
    assert any(s.body.startswith("31") for s in dsk_verses)
    # The reference is "Mark 6:30-32 (MSG)" but the quote starts at 31, so the
    # range must not be left stranded at the head of the body.
    assert not any(s.body.startswith("30-32") for s in dsk_verses)
    assert all("(MSG)" not in s.body for s in dsk_verses)


def test_verse_after_point_masters(tmp_path):
    """[VERSE-AFTER-POINT] is what builds the point-plus-verse slide."""
    variant = verse_after_point_variant(OUTLINES / "Sermon BC.docx", tmp_path / "BC_VAP.docx")
    lw, dsk, _ = map_slides(parse_outline(variant))

    truthful_pre = next(
        s for s in lw if s.master == "NUMBERED POINT PRE" and "Truthful" in s.text_items.get(1, "")
    )
    assert truthful_pre.transition is not None
    assert truthful_pre.transition.effect == "magic_move"
    assert truthful_pre.transition.duration == 1
    assert truthful_pre.transition.match == "word"

    truthful_post = next(s for s in lw if s.master == "NUMBERED POINT POST")
    assert truthful_post.transition is None
    assert "27" in truthful_post.body or any(
        "27" in "".join(r.text for r in runs) for runs in truthful_post.styled_items.values()
    )
    post_title = truthful_post.text_items.get(1, "")
    assert "Truthful" in post_title
    assert "\n" in post_title
    # The POST carries its own cue now, so its operator tag sits there.
    assert truthful_post.bind == "cue"

    # Examine is followed by another point, so it neither moves nor gains a POST.
    examine = next(
        s for s in lw if s.master == "NON-NUMBERED POINT PRE" and "Examine" in (s.body or "")
    )
    assert examine.transition is None
    assert not any(
        s.master == "NON-NUMBERED POINT POST" and "Examine" in (s.body or "") for s in lw
    )

    communion_post = next(
        s for s in lw if s.master == "NON-NUMBERED POINT POST" and "Communion" in (s.body or "")
    )
    post_text = communion_post.text_items.get(1, "")
    assert post_text.count("\n") == 2
    assert "Powerful Presence" in post_text
    assert not (communion_post.text_items.get(2) or "").strip()
    assert not (communion_post.text_items.get(3) or "").strip()
    # Too long for the lower-third column, so DSK skips this one.
    assert not any(s.role == "post" and "Communion" in (s.body or "") for s in dsk)

    dsk_pre = next(s for s in dsk if s.master == "Num Point with Verse-Pre")
    assert dsk_pre.transition is not None
    assert dsk_pre.transition.match == "word"
    dsk_post = next(s for s in dsk if s.master == "Num Point with Verse-Post")
    extra_text = " ".join(e.get("text") or "" for e in dsk_post.extra_text_items)
    assert "Truthful" in extra_text
    assert any(e.get("text") == "1" for e in dsk_post.extra_text_items)
    verse_box = dsk_post.text_items.get(1, "")
    header_box = dsk_post.text_items.get(2, "")
    assert "27" in verse_box or any(
        "27" in "".join(r.text for r in runs) for runs in dsk_post.styled_items.values()
    )
    assert "Ezekiel" in header_box
    assert "Spirit" not in header_box
    assert "Truthful" not in verse_box
    assert 3 not in dsk_post.text_items


def test_orphan_verse_after_point_is_flagged(tmp_path):
    path = build_outline(
        tmp_path / "orphan.docx",
        ["[VERSE-AFTER-POINT] Ezekiel 36:26 I will give you a new heart."],
    )
    lw, _dsk, flags = map_slides(parse_outline(path))
    assert any("no point cue before it" in f.message for f in flags)
    # Falls back to a plain verse rather than dropping the content.
    assert any(s.master == "VERSES" for s in lw)
    assert not any(s.role == "post" for s in lw)


@pytest.mark.parametrize(
    "name", ["Sermon BC.docx", "Offering JX.docx"]
)
def test_every_slide_has_exactly_one_cue(tmp_path, name):
    """The invariant the outline checker relies on: one cue, one slide."""
    outline = parse_outline(OUTLINES / name)
    lw, dsk, _ = map_slides(outline)
    cued = annotate_outline(outline, lw, dsk, tmp_path / f"{Path(name).stem}_CUED.docx")
    assert operator_cue_counts(cued) == (len(lw), len(dsk))


def test_every_slide_has_exactly_one_cue_with_verse_after_point(tmp_path):
    variant = verse_after_point_variant(OUTLINES / "Sermon BC.docx", tmp_path / "BC_VAP.docx")
    outline = parse_outline(variant)
    lw, dsk, _ = map_slides(outline)
    cued = annotate_outline(outline, lw, dsk, tmp_path / "BC_VAP_CUED.docx")
    assert any(s.role == "post" for s in lw), "the variant should exercise POST slides"
    assert operator_cue_counts(cued) == (len(lw), len(dsk))


def test_deprecated_verse_from_previous_alias(tmp_path):
    """The retired spelling still maps, but says so."""
    path = build_outline(
        tmp_path / "deprecated.docx",
        [
            "[VERSE] Ezekiel 36:26 I will give you a new heart and a new spirit.",
            "[VERSE-FROM-PREVIOUS] and put a new spirit in you.",
        ],
    )
    outline = parse_outline(path)
    assert [b.cue_tag for b in outline.blocks] == ["VERSE", "VERSE-CONTINUED"]
    _lw, _dsk, flags = map_slides(outline)
    stale = [f for f in flags if f.rule == "cue.deprecated_alias"]
    assert len(stale) == 1
    assert "[VERSE-CONTINUED]" in stale[0].message


def test_ref_tail_never_eats_a_verse_number():
    """An inline reference leaves "30-32 (MSG)" behind; a lone number is a verse.

    `_resolve_ref` returns only "Book Chapter", so dropping len(ref) strips
    "Mark 6" and leaves ":30-32 (MSG) 31 …" (the colon is gone by then). The
    tail therefore has to require a range or a translation: matching a bare
    number would delete the verse number itself.
    """
    from obed_edom.slide_map import REF_TAIL_RE, VERSE_LEAD_RE

    def dropped(text: str) -> str | None:
        match = REF_TAIL_RE.match(text)
        if match and VERSE_LEAD_RE.match(match.group("rest")):
            return text[match.end("tail") :].strip()
        return None

    assert dropped("30-32 (MSG) 31 “If God gives") == "31 “If God gives"
    assert dropped("30-32 31 “If God gives") == "31 “If God gives"
    assert dropped("6 (MSG) 31 “If God") == "31 “If God"
    # A bare number is a verse number, not a reference range.
    assert dropped("26 I will give you a new heart") is None
    assert dropped("26 27 And I will put my Spirit") is None
    assert dropped("31 “If God gives such attention") is None


def test_list_number_resolver():
    from docx import Document
    from obed_edom.parse_outline import ListNumberResolver

    doc = Document(str(OUTLINES / "Sermon BC.docx"))
    resolver = ListNumberResolver(doc)
    truthful_para = next(p for p in doc.paragraphs if "Be Truthful in Love" in p.text)
    number = resolver.number_for_paragraph(truthful_para)
    outline = parse_outline(OUTLINES / "Sermon BC.docx")
    draft = next(d for d in outline.blocks if "Truthful" in d.body)
    assert draft.point_number == 1
    if number is not None:
        assert number == 1


def test_point_styled_runs_and_fit():
    from obed_edom.models import SlideDraft, TextSpan
    from obed_edom.slide_map import _fit_point_runs, _point_styled_runs

    draft = SlideDraft(
        cue_tag="NUM-POINT",
        body_spans=[
            TextSpan(text="Be", bold=False),
            TextSpan(text=" Truthful ", bold=True),
            TextSpan(text="in Love", bold=False),
        ],
    )
    runs = _point_styled_runs(draft)
    assert [r.style for r in runs] == ["normal", "highlight", "normal"]
    assert "Truthful" in "".join(r.text for r in runs if r.style == "highlight")

    long_runs = [StyledRun(text="word " * 40, style="normal")]
    lines, size = _fit_point_runs(
        long_runs, box_width=820, base_size=180, min_size=64, em=0.62, max_lines=3
    )
    assert len(lines) >= 2
    assert size <= 180
    short_runs = [StyledRun(text="Be Truthful in Love", style="highlight")]
    one, pre_size = _fit_point_runs(
        short_runs, box_width=1680, base_size=180, min_size=64, em=0.5, max_lines=3
    )
    assert len(one) == 1
    assert pre_size >= 170
    from obed_edom.slide_map import _points_text_items

    mapping = {"line1": 1, "line2": 2, "line3": 3, "point_number": 4}
    items = _points_text_items(mapping, ["Be Truthful in Love"], point_number=9)
    assert items == {1: "Be Truthful in Love", 2: "", 3: "", 4: "9"}


def test_split_styled_runs():
    runs = [
        StyledRun(text="31", style="verse_number"),
        StyledRun(text=" " + ("word " * 50), style="normal"),
        StyledRun(text="highlight", style="highlight"),
    ]
    chunks = _split_styled_runs(runs, 80)
    assert len(chunks) >= 2
    assert chunks[0][0].style == "verse_number"
    assert chunks[0][0].text == "31"
    assert all(sum(len(r.text) for r in chunk) <= 120 for chunk in chunks)
    assert min(sum(len(r.text) for r in chunk) for chunk in chunks) >= 8
    assert any(r.style == "highlight" for chunk in chunks for r in chunk)
    long_body = "word " * 80
    long_runs = [StyledRun(text="31", style="verse_number"), StyledRun(text=" " + long_body, style="normal")]
    long_chunks = _split_styled_runs(long_runs, 190)
    assert len(long_chunks) >= 2
    assert min(sum(len(r.text) for r in chunk) for chunk in chunks) >= 40


def test_prepare_styled_runs():
    from obed_edom.keynote import _prepare_styled_runs

    first, rest, prepared = _prepare_styled_runs(
        [
            {"text": "33", "style": "verse_number"},
            {"text": " People", "style": "normal"},
            {"text": "highlight", "style": "highlight"},
        ]
    )
    assert first == "33"
    assert rest == " Peoplehighlight"
    assert [p["style"] for p in prepared] == ["verse_number", "normal", "highlight"]

    first, rest, prepared = _prepare_styled_runs(
        [
            {"text": "31", "style": "verse_number"},
            {"text": " one ", "style": "normal"},
            {"text": "32", "style": "verse_number"},
            {"text": " two", "style": "normal"},
        ]
    )
    assert first == "31"
    assert rest == " one 32 two"
    assert prepared[2]["text"] == "32"

    from obed_edom.keynote import _later_verse_jobs

    later = _later_verse_jobs(
        [
            {"text": "26", "style": "verse_number"},
            {"text": " a ", "style": "normal"},
            {"text": "27", "style": "verse_number"},
            {"text": " Drink from it, all of you.", "style": "normal"},
            {"text": "28", "style": "verse_number"},
            {"text": " This is My Blood", "style": "normal"},
        ]
    )
    assert [v["digits"] for v in later] == ["27", "28"]
    assert [v["start"] for v in later] == [6, 35]
    # Later verses are found through the copy that follows them, not a placeholder token.
    assert later[0]["anchor"] == " Drink from it, all of y"
    assert later[1]["anchor"] == " This is My Blood"

    first, rest, prepared = _prepare_styled_runs(
        [
            {"text": " ", "style": "normal"},
            {"text": "26", "style": "verse_number"},
            {"text": " Take and eat", "style": "normal"},
        ]
    )
    assert first == "26"
    assert rest.startswith(" Take and eat")
    assert prepared[0]["style"] == "verse_number"


def test_point_applescript_replaces_placeholder(tmp_path):
    """Point titles must replace the whole text box, not seed into Bold/Normal."""
    from pathlib import Path

    from obed_edom.keynote import STYLE_PALETTES, _build_applescript, _plan_payload, _prepare_styled_runs

    assert STYLE_PALETTES["lw"]["verse_number"]["size"] == 70

    # POST slides need a cue that asks for them, and the Matthew passage is
    # wanted both beside its point and on its own.
    variant = verse_after_point_variant(OUTLINES / "Sermon BC.docx", tmp_path / "BC_VAP.docx")
    duplicate_as_plain_verse(variant, "Take and eat")
    sermon = parse_outline(variant)
    lw, dsk, _ = map_slides(sermon)
    pre = next(s for s in lw if s.master == "NUMBERED POINT PRE" and "Truthful" in s.body)
    post = next(s for s in lw if s.master == "NUMBERED POINT POST")
    script = _build_applescript(_plan_payload([pre, post], Path("/tmp/x.key"), None))
    assert 'set object text of text item 1 to "Be Truthful in Love "' in script
    # POST verse keeps the template superscript seed; point titles do not.
    assert "set character 1 to" in script
    assert 'set character 1 to "Be Truthful' not in script
    assert "set size of characters 1 thru 20 to 45" not in script
    assert "46.67" not in script
    assert "delete characters 1 thru clearN" in script

    verse = next(s for s in lw if s.role == "verse" and "new heart" in (s.body or ""))
    verse_script = _build_applescript(_plan_payload([verse], Path("/tmp/x.key"), None))
    assert "46.67" not in verse_script
    assert "baseline offset" not in verse_script
    assert "usedSeed" in verse_script
    assert "set character 1 to" in verse_script
    assert "to 70" in verse_script

    dsk_pre = next(s for s in dsk if s.master == "Num Point with Verse-Pre")
    dsk_post = next(s for s in dsk if s.master == "Num Point with Verse-Post")
    dsk_script = _build_applescript(_plan_payload([dsk_pre, dsk_post], Path("/tmp/dsk.key"), None))
    assert "duplicate slide magicDonor" in dsk_script
    assert "make new text item with properties" in dsk_script
    assert "usedSeed" in dsk_script
    assert "set object text of text item 2 to" in dsk_script
    assert "Ezekiel 36" in dsk_script

    from obed_edom.slide_map import _styled_verse_runs

    matt = next(b for b in sermon.blocks if "Take and eat" in b.body)
    matt_runs = _styled_verse_runs(matt, "Matthew 26")
    assert matt_runs[0].style == "verse_number"
    assert matt_runs[0].text.strip() == "26"
    first, rest, _ = _prepare_styled_runs(
        [{"text": r.text, "style": r.style} for r in matt_runs]
    )
    assert first == "26"
    assert "Take and eat" in rest
    assert "27" in rest
    assert "28" in rest
    assert "²" not in rest
    assert "⁷" not in rest
    assert "⁸" not in rest
    later_numbers = [
        r.text for r in matt_runs if r.style == "verse_number" and r.text.strip() != "26"
    ]
    assert "27" in later_numbers
    assert "28" in later_numbers

    matt_slide = next(s for s in lw if s.role == "verse" and "Take and eat" in (s.body or ""))
    matt_script = _build_applescript(_plan_payload([matt_slide], Path("/tmp/matt.key"), None))
    assert 'click menu item "Superscript"' not in matt_script
    assert "System Events" not in matt_script
    assert "offset of" not in matt_script
    assert "set seedBox to duplicate mainBox" not in matt_script
    assert "‡" not in matt_script
    assert "†" not in matt_script
    assert "27" in matt_script
    assert "28" in matt_script
    assert f'using terms from application id "{keynote_app.bundle_id()}"' in matt_script
    assert "to (size of character 1)" not in matt_script
    assert "²" not in matt_script
    assert "⁷" not in matt_script
    assert "⁸" not in matt_script
    assert "set font of characters 3 thru 91" not in matt_script
    assert 'set font of characters 37 thru 38 to "AzoSans-Regular"' not in matt_script
    assert 'set font of characters 69 thru 70 to "AzoSans-Regular"' not in matt_script

    assert any("AzoSans-Bold" in ln for ln in matt_script.splitlines())


def test_later_verse_numbers_get_the_template_character_style():
    """Pass 2 must stay GUI-driven and must carry the template's character style.

    Verified against Keynote 14.5 (see the "Later verse numbers" section of
    .agents/skills/obed-edom/SKILL.md). Keynote's AppleScript dictionary has
    no style support at all, and superscript is not a character property, so the
    deck's own verse-number style can only be applied through the UI. Do not
    "simplify" this into a pure-AppleScript pass: every scriptable route asserted
    against below was tried and rejected by Keynote itself, each failing silently
    inside a ``try`` block.
    """
    from obed_edom.keynote import (
        _build_superscript_fix_script,
        _collect_superscript_jobs,
        _read_superscript_report,
    )

    sermon = parse_outline(OUTLINES / "Sermon BC.docx")
    lw, _, _ = map_slides(sermon)
    matt_slide = next(s for s in lw if s.role == "verse" and "Take and eat" in (s.body or ""))

    jobs = _collect_superscript_jobs([matt_slide])
    assert jobs and jobs[0]["laterVerses"]
    # Find is the only scripted way to place a selection, so every verse number
    # needs the copy that follows it as a search anchor -- the seed included,
    # since its style is what the later ones are copied from.
    assert all(v["anchor"] for v in jobs[0]["laterVerses"])
    assert jobs[0]["seed"]["anchor"]
    assert jobs[0]["seed"]["digits"] == "26"

    fix_script = _build_superscript_fix_script(Path("/tmp/matt.key"), jobs)
    # The style comes from the deck's own first verse number, so no style name is
    # hardcoded: LW's template calls it "SuperScript", DSK's "Verse Number".
    assert fix_script.count('click menu item "Copy Style"') == 1
    assert 'click menu item "Paste Style"' in fix_script
    assert "SuperScript" not in fix_script
    assert "Verse Number" not in fix_script
    # Raw baseline formatting would bypass the character style.
    assert "Baseline" not in fix_script
    assert 'menu item "Superscript"' not in fix_script
    assert 'keystroke "f" using {command down}' in fix_script
    assert "key code 123 using {shift down}" in fix_script

    # "Shapes can not be copied" / "Words can not be copied".
    assert "duplicate" not in fix_script
    # Copies text only, never the superscript attribute.
    assert "to character 1 of" not in fix_script
    assert "set character 37 to character 1" not in fix_script
    # `size` is the base size that superscript renders at 2/3; it cannot fake one.
    assert "46.6" not in fix_script
    # Mixing Latin-1 ²³ with the superscripts block ⁴-⁹ gives mismatched digits.
    for glyph in "²³⁴⁵⁶⁷⁸⁹":
        assert glyph not in fix_script

    # A denied Accessibility grant has to surface, not pass silently.
    denied = _read_superscript_report(
        "s=7 ti=1 c1=30.0 c37=45.0 c69=45.0 gui= [-1743: not authorized] exported=false"
    )
    assert denied["accessibilityDenied"] is True
    assert denied["allSuperscript"] is False
    assert denied["exported"] is False

    # Every text box in the report is checked, including the first one.
    applied = _read_superscript_report(
        "s=8 ti=5 c1=46.67 c37=46.67 c69=46.67 s=9 ti=2 c1=46.67 c37=46.67 c69=46.67"
        " gui= exported=true"
    )
    assert len(applied["boxes"]) == 2
    assert applied["allSuperscript"] is True
    assert applied["exported"] is True
    unfixed = _read_superscript_report("s=8 ti=5 c1=46.67 c37=70.0 c69=70.0 gui= exported=true")
    assert unfixed["allSuperscript"] is False


def test_repeated_verse_box_is_styled_on_every_slide(tmp_path):
    """Find cycles through matches, so each anchor is applied once per occurrence.

    Cueing a passage as both [VERSE-AFTER-POINT] and [VERSE] puts the same verse
    box on two slides. Applying an anchor once styles a single instance and
    silently leaves the other on the baseline.
    """
    from obed_edom.keynote import (
        _build_superscript_fix_script,
        _collect_superscript_jobs,
        _superscript_anchor_plan,
    )

    variant = verse_after_point_variant(OUTLINES / "Sermon BC.docx", tmp_path / "BC_VAP.docx")
    duplicate_as_plain_verse(variant, "Take and eat")
    lw, _, _ = map_slides(parse_outline(variant))
    jobs = _collect_superscript_jobs(lw)
    matt_jobs = [j for j in jobs if "Take and eat" in j["marker"]]
    assert len(matt_jobs) == 2, "the Matthew verse box should appear on two slides"

    plan = _superscript_anchor_plan(jobs)
    assert plan
    for entry in plan:
        assert entry["occurrences"] == 2

    script = _build_superscript_fix_script(Path("/tmp/matt.key"), jobs)
    assert "repeat 2 times" in script
    # One copy of the seed style, then a paste per occurrence of each anchor.
    assert script.count('click menu item "Copy Style"') == 1
    assert script.count('click menu item "Paste Style"') == len(plan)


def test_pass_one_hands_the_open_deck_to_pass_two():
    """Pass 1 leaves the deck open and defers its export when pass 2 follows.

    Closing and reopening the same file was pure overhead, and exporting in pass 1
    rendered every slide twice with the verse numbers still on the baseline.
    """
    from obed_edom.keynote import _build_applescript, _plan_payload

    lw, _, _ = map_slides(parse_outline(OUTLINES / "Sermon BC.docx"))
    matt_slide = next(s for s in lw if s.role == "verse" and "Take and eat" in (s.body or ""))

    plan = _plan_payload([matt_slide], Path("/tmp/matt.key"), Path("/tmp/shots"))
    assert plan["superscriptJobs"], "this slide should need pass 2"

    plan["superscriptFix"] = True
    handed_off = _build_applescript(plan)
    assert "close theDoc" not in handed_off
    assert 'set exported to "deferred"' in handed_off
    assert "export theDoc" not in handed_off

    # Handing off is only safe while pass 2 has work it can do, so every job must
    # carry the anchors pass 2 needs; otherwise the deck is left open forever.
    from obed_edom.keynote import _build_superscript_fix_script

    for job in plan["superscriptJobs"]:
        assert job["seed"]["anchor"]
        assert job["laterVerses"]
        assert all(v["anchor"] for v in job["laterVerses"])
    assert _build_superscript_fix_script(Path("/tmp/matt.key"), plan["superscriptJobs"])

    # With no pass 2 to follow, pass 1 stays responsible for export and close.
    plan["superscriptFix"] = False
    standalone = _build_applescript(plan)
    assert "close theDoc" in standalone
    assert "export theDoc" in standalone
    assert 'set exported to "deferred"' not in standalone


def test_review_pdf_and_slide_kinds():
    from obed_edom.pipeline import generate
    from obed_edom.report import slide_kind

    offering = parse_outline(OUTLINES / "Offering JX.docx")
    lw, dsk, _ = map_slides(offering)
    assert any(slide_kind(s) == "Giving options (paste the graphic)" for s in lw)
    assert any(slide_kind(s) == "Ways to give (QR code)" for s in dsk)
    assert any(slide_kind(s) == "Bible verse" for s in dsk)

    with patch("obed_edom.bible.fetch_passage", return_value=(None, "mocked")):
        result = generate(OUTLINES / "Sermon BC.docx", make_keynote=False, check_visuals=False)
    assert result.review_path.suffix == ".pdf"
    assert result.review_path.exists()
    header = result.review_path.read_bytes()[:5]
    assert header == b"%PDF-"
    pdf = result.review_path.read_bytes()
    assert b"Slide check" not in pdf
    assert b"What's on it" not in pdf
    assert b"This folder has two Keynote files" not in pdf
    assert b"Passion for God" in pdf
    assert b"Truth and Contentment" in pdf
    assert result.cued_docx is not None and result.cued_docx.exists()
    assert not (result.output_dir / "review.md").exists()
    assert not list(result.output_dir.glob("*.json"))


def test_passage_header():
    from obed_edom.slide_map import _passage_header

    assert _passage_header("Ezekiel 36", "NIV", "Truth and Contentment") == "Ezekiel 36 • Truth and Contentment"
    assert _passage_header("Matthew 6", "MSG", "") == "Matthew 6 (MSG)"
    assert _passage_header("", "NIV", "Series Title") == "Series Title"


def test_annotate_offering_splits(tmp_path):
    outline = parse_outline(OUTLINES / "Offering JX.docx")
    lw, dsk, _ = map_slides(outline)
    dest = tmp_path / "test_cued_offering.docx"
    annotate_outline(outline, lw, dsk, dest)
    tags = extract_operator_cues(dest)
    assert tags[0] == "LW-OFFERING FILLER"
    assert tags[1] == "DSK-PP-QR CODE"
    assert "LW-GIVING OPTIONS" in tags
    assert "DSK-PP-GIVING OPTIONS" in tags
    assert tags.count("LW-OFFERING FILLER") == 2
    assert tags.count("DSK-PP-QR CODE") == 2

    from docx import Document

    doc = Document(str(dest))
    verse_para = next(p for p in doc.paragraphs if "31" in p.text and "wildflowers" in p.text)
    text = verse_para.text
    assert text.index("[DSK-PP]") < text.index("31")
    assert text.index("31") < text.index("[DSK-PP]", text.index("31"))
    second = text.index("[DSK-PP]", text.index("31"))
    assert second < text.index("32")
    assert text.index("32") < text.rindex("[DSK-PP]")
    assert text.rindex("[DSK-PP]") < text.index("33")
    expected = extract_operator_cues(OUTLINES / "Offering JX_EXPECTED_RESULT.docx")
    assert expected[0] == "LW-OFFERING FILLER"
    assert "DSK-PP-GIVING OPTIONS" in expected


def test_annotate_sermon_point_and_title(tmp_path):
    outline = parse_outline(OUTLINES / "Sermon BC.docx")
    lw, dsk, _ = map_slides(outline)
    dest = tmp_path / "test_cued_sermon.docx"
    annotate_outline(outline, lw, dsk, dest)
    tags = extract_operator_cues(dest)
    assert tags[0] == "LW-TITLE"
    assert "LW" in tags
    assert "DSK-PP" in tags
    from docx import Document

    doc = Document(str(dest))
    title_para = next(p for p in doc.paragraphs if "Following" in p.text and "Jesus" in p.text)
    assert "[LW-TITLE]" in title_para.text
    assert "[TITLE]" not in title_para.text
    truthful = next(p for p in doc.paragraphs if "Be Truthful in Love" in p.text)
    assert "[LW]" in truthful.text
    assert "[DSK-PP]" in truthful.text
    assert "[NUM-POINT]" not in truthful.text

    matt = next(p for p in doc.paragraphs if "Take and eat" in p.text and "My Blood" in p.text)
    compact = matt.text.replace(" ", "")
    assert "[LW][DSK-PP][LW][DSK-PP]" not in compact
    body = matt.text[matt.text.index("[LW]") :]
    assert body.index("[LW]") < body.index("26")
    assert body.index("[DSK-PP]") < body.index("26")
    assert body.count("[LW]") == 1
    assert body.count("[DSK-PP]") == 1

    def _run_highlight(run) -> str | None:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            return None
        hl = rpr.find(qn("w:highlight"))
        if hl is None:
            return None
        return hl.get(qn("w:val"))

    lw_highlights = {
        _run_highlight(run)
        for para in doc.paragraphs
        for run in para.runs
        if (run.text or "").startswith("[LW")
    }
    assert lw_highlights == {"cyan"}
    dsk_highlights = {
        _run_highlight(run)
        for para in doc.paragraphs
        for run in para.runs
        if (run.text or "").startswith("[DSK")
    }
    assert dsk_highlights == {"yellow"}


def test_annotate_verse_cues_at_chunk_starts(tmp_path):
    """If 26–27 fit on slide 1 and 28 starts slide 2, cues sit before 26 and before 28."""
    outline = parse_outline(OUTLINES / "Sermon BC.docx")
    lw, dsk, _ = map_slides(outline)
    matt_block = next(i for i, b in enumerate(outline.blocks) if "Take and eat" in b.body)
    lw_verse = next(s for s in lw if s.role == "verse" and s.block_index == matt_block)
    dsk_verse = next(s for s in dsk if s.role == "verse" and s.block_index == matt_block)
    extra_lw = replace(
        lw_verse,
        chunk_index=1,
        anchor_verse="28",
        body="28 … This is My Blood …",
    )
    extra_dsk = replace(
        dsk_verse,
        chunk_index=1,
        anchor_verse="28",
        body="28 … This is My Blood …",
    )
    dest = tmp_path / "test_cued_sermon_split.docx"
    annotate_outline(outline, lw + [extra_lw], dsk + [extra_dsk], dest)
    from docx import Document

    doc = Document(str(dest))
    matt = next(p for p in doc.paragraphs if "Take and eat" in p.text and "My Blood" in p.text)
    text = matt.text[matt.text.index("[LW]") :]
    i26 = text.index("26")
    i28 = text.index("28", i26)
    first_lw = text.index("[LW]")
    first_dsk = text.index("[DSK-PP]")
    assert first_lw < i26
    assert first_dsk < i26
    second_lw = text.index("[LW]", i26)
    second_dsk = text.index("[DSK-PP]", i26)
    assert i26 < second_lw < i28
    assert i26 < second_dsk < i28
    assert text.count("[LW]") == 2
    assert text.count("[DSK-PP]") == 2
    assert "[LW][DSK-PP][LW][DSK-PP]" not in text.replace(" ", "")


GATEWAY_HTML = """
<html><body>
<div class="passage-content passage-class">
<p><sup class="versenum">26</sup> I will give you a new heart and put a new spirit in you;
<sup class="versenum">27</sup> And I will put my Spirit in you and move you to follow my decrees.</p>
</div></div>
</body></html>
"""

GATEWAY_HTML_CHROME = """
<html><body>
<div>Bible Gateway logo Advanced Search Available Versions Bible Gateway Plus Log In/Sign Up</div>
<div class='passage-content passage-class-0'><div class="version-MSG result-text-style-normal text-html">
<p><span class="text Mark-6-31"><sup class="versenum">31</sup>The apostles then rendezvoused with Jesus.</span></p>
<a class="full-chap-link" href="/passage/?search=Mark%206&version=MSG">Read full chapter</a>
</div></div>
<div class="copyright-table"><div class="publisher-info-bottom">The Message</div></div>
</body></html>
"""


def test_bible_gateway_parse_and_fetch():
    text = _parse_gateway_html(GATEWAY_HTML)
    assert "26" in text
    assert "new heart" in text
    assert "27" in text

    chrome = _parse_gateway_html(GATEWAY_HTML_CHROME)
    assert "rendezvoused" in chrome.lower()
    assert "available versions" not in chrome.lower()
    assert "bible gateway logo" not in chrome.lower()

    class _Resp:
        ok = True
        status_code = 200
        text = GATEWAY_HTML

    from obed_edom import bible as bible_mod

    bible_mod._GATEWAY_CACHE.clear()
    with patch("obed_edom.bible.requests.get", return_value=_Resp()):
        official, source = fetch_passage("Ezekiel", 36, 26, 27, "NIV")
    assert official is not None
    assert "Bible Gateway NIV" in source
    assert "new heart" in official.lower()


def test_wrong_gospel_citation_is_flagged():
    from obed_edom.report import _action_items, _bible_notes

    outline = parse_outline(OUTLINES / "Offering JX.docx")
    matt = (
        "If God gives such attention to the appearance of wildflowers most of which "
        "are never even seen don't you think he'll attend to you take pride in you "
        "do his best for you What I'm trying to do here is to get you to relax"
    )
    mark = "The apostles then rendezvoused with Jesus and reported on all that they had done and taught."
    luke = "Give away your life and you will find life given back."
    john = "The next day the crowd that had stayed on the opposite shore."

    def fake_fetch(book, chapter, verse, verse_end, translation):
        texts = {"Matthew": matt, "Mark": mark, "Luke": luke, "John": john}
        return (texts.get(book, "unrelated"), f"Bible Gateway {translation}")

    with patch("obed_edom.bible.fetch_passage", side_effect=fake_fetch):
        flags = check_bible(outline)
    mix = [
        f
        for f in flags
        if f.severity == "error" and "Cited as" in f.message and "Matthew" in f.message
    ]
    assert mix, [f.message for f in flags if f.category == "bible"]
    notes = _bible_notes(flags)
    assert any("Matthew" in n and "Cited as" in n for n in notes)
    actions = _action_items(flags)
    assert any("Matthew" in a for a in actions)


def test_verse_continued_cue_and_full_verse_on_second_slide():
    from obed_edom.parse_outline import normalize_cue
    from obed_edom.models import SlideSpec
    from obed_edom.keynote import _append_seeded_text, _build_applescript, _plan_payload

    cue = normalize_cue("[VERSE-CONTINUED]")
    assert cue.tag == "VERSE-CONTINUED"
    assert normalize_cue("[VERSE-FROM-PREVIOUS]").tag == "VERSE-CONTINUED"

    offering = parse_outline(OUTLINES / "Offering JX.docx")
    lw, dsk, _ = map_slides(offering)
    lw_full = [s for s in lw if s.is_verse and "Steep your life" in s.body]
    assert len(lw_full) == 1
    assert lw_full[0].semantic_tag == "VERSE-CONTINUED"
    assert lw_full[0].body.strip().startswith("33")
    assert "People who" in lw_full[0].body
    assert "but you know both God" in lw_full[0].body
    runs = list(lw_full[0].styled_items.values())[0]
    assert runs[0].style == "verse_number"
    first_33 = next(
        s
        for s in lw
        if s.is_verse and s.body.strip().startswith("33") and "Steep" not in s.body
    )
    assert "but you know both God" in first_33.body

    dsk_full = [s for s in dsk if s.is_verse and "Steep your life" in s.body]
    assert len(dsk_full) == 1
    assert "People who" in dsk_full[0].body
    assert dsk_full[0].body.strip().startswith("33")

    lines = []
    _append_seeded_text(lines, 2, "", "Steep your life", mode="body_only")
    script = "\n".join(lines)
    assert "delete characters 1 thru (bodyIdx - 1)" in script
    assert "usedSeed" in script

    fragment = SlideSpec(
        deck="lw",
        cue_tag="LW",
        master="VERSES",
        body="Steep your life in God-reality",
        text_items={2: "Steep your life in God-reality"},
        styled_items={2: [StyledRun(text="Steep your life in God-reality", style="highlight")]},
        is_verse=True,
        semantic_tag="VERSE-CONTINUED",
    )
    frag_script = _build_applescript(_plan_payload([fragment], Path("/tmp/cont.key"), None))
    assert "delete characters 1 thru (bodyIdx - 1)" in frag_script
    assert "usedSeed" in frag_script
    used = frag_script.index("set usedSeed to false")
    replace = frag_script.index(
        'set object text of text item 2 to "Steep your life in God-reality"'
    )
    assert "if usedSeed then" in frag_script[used:replace]


if __name__ == "__main__":
    test_sermon_cues()
    test_offering_cues()
    test_mapping_masters()
    test_list_number_resolver()
    test_point_styled_runs_and_fit()
    test_split_styled_runs()
    test_prepare_styled_runs()
    test_point_applescript_replaces_placeholder()
    test_passage_header()
    test_bible_gateway_parse_and_fetch()
    test_wrong_gospel_citation_is_flagged()
    test_verse_continued_cue_and_full_verse_on_second_slide()
    test_annotate_offering_splits()
    test_annotate_sermon_point_and_title()
    test_annotate_verse_cues_at_chunk_starts()
    test_review_pdf_and_slide_kinds()
    print("offline tests passed")
